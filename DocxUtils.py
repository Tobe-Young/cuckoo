# -*- coding: utf-8 -*-
from docx import Document
from typing import Dict, List

class KeyChanger:
    def __init__(self, p, key, value) -> None:
        self.p = p
        self.key = key
        self.value = value
        self.run_text = ""
        self.runs_indexes: List = []
        self.run_char_indexes: List = []
        self.runs_to_change: Dict = {}

    def _initialize(self) -> None:
        run_index = 0
        for run in self.p.runs:
            self.run_text += run.text
            self.runs_indexes += [run_index for _ in run.text]
            self.run_char_indexes += [char_index for char_index, char in enumerate(run.text)]
            run_index += 1

    def replace(self) -> None:
        self._initialize()
        parsed_key_length = len(self.key)
        index_to_replace = self.run_text.find(self.key)

        for i in range(parsed_key_length):
            index = index_to_replace + i
            run_index = self.runs_indexes[index]
            run = self.p.runs[run_index]
            run_char_index = self.run_char_indexes[index]

            if not self.runs_to_change.get(run_index):
                self.runs_to_change[run_index] = [char for char_index, char in enumerate(run.text)]

            run_to_change: Dict = self.runs_to_change.get(run_index)  # type: ignore[assignment]
            if index == index_to_replace:
                run_to_change[run_char_index] = self.value
            else:
                run_to_change[run_char_index] = ""

        # make the real replace
        for index, text in self.runs_to_change.items():
            run = self.p.runs[index]
            run.text = "".join(text)
class DocxHelper:
    def __init__(self, docx_path, make_copy = False):
        try:
            self.doc = Document(docx_path)
        except Exception as e:
            print(str(e))

        self.doc_path = docx_path
        self.make_copy = make_copy

    def save(self):
        if self.make_copy:
            formatter = "{}_Copy.docx"
            str_path = str(self.doc_path)
            start = str_path.find(".docx")
            self.doc.save(formatter.format(str_path[0: int(start)]))
        else:    
            self.doc.save(self.doc_path)

    def is_string_exists(self, find_text):
        for paragraph in self.doc.paragraphs:
            if find_text in paragraph.text:
                return True
        
        for table in self.doc.tables:
            for row in table.rows:  
                for cell in row.cells:
                    if find_text in cell.text:
                        return True
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            return True
        return False      
    
    def find_replace(self, find_text, replace_text):
        for paragraph in self.doc.paragraphs:
            self.replace_in_paragrah(paragraph, find_text, replace_text)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.find_replace_cell(cell, find_text, replace_text)

    def replace_in_paragrah(self, paragraph, search_text, replace_text):
        if search_text not in paragraph.text:
            return
        replacer = KeyChanger(paragraph, search_text, replace_text)
        replacer.replace()

    def find_paragrahs(self, find_text):
        res = []
        for paragraph in self.doc.paragraphs:
            if find_text in paragraph.text:
                res.append(paragraph.text)

        for table in self.doc.tables:
            for row in table.rows:  
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            res.append(paragraph.text)
        return res
    
    def async_find_replace(self, find_text, replace_text):
        self.find_replace(find_text, replace_text)
                    
    def find_replace_cell(self, cell, find_text, replace_text):
        for paragraph in cell.paragraphs:
            self.replace_in_paragrah(paragraph, find_text, replace_text)
    